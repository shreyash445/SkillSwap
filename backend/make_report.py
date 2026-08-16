from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def h(text, size=16, color=(205, 255, 87)):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(*color)
    return p

def para(text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
    p.add_run(text)
    return p

def bullet(text):
    doc.add_paragraph(text, style="List Bullet")

# Title
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("SkillSwap MVP — Progress Report")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = RGBColor(0, 0, 0)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Status: Significant UI improvements with a light backend layer")
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(90, 90, 90)

doc.add_paragraph()

h("1. Summary", 15, (20, 20, 20))
para(
    "This sprint delivered a major upgrade to the SkillSwap MVP mobile experience while adding a "
    "small amount of backend work. The app now feels polished, modern, and production-ready from a "
    "visual standpoint, with the backend extended only where the UI required it."
)

doc.add_paragraph()
h("2. UI Improvements (Significant)", 15, (20, 20, 20))
bullet("Full-black dark theme — removed all dark-blue tones in favour of pure black and neutral greys.")
bullet("Animated logo splash — a smooth spring animation plays on launch before routing the user to onboarding or home.")
bullet("Glassmorphism dock — the bottom navigation is now a floating, frosted-glass capsule dock instead of a plain bar.")
bullet("Soft, subtle spring animations on bottom sheets and sub-menus for a calmer feel.")
bullet("Keyboard handling fixed — the on-screen keyboard now pushes content up instead of covering the input fields.")
bullet("Custom designed components, typography (Sora / Instrument Sans / Space Mono), icons, and badges.")
bullet("Built as a native mobile app with Expo SDK 54 and React Native, running in Expo Go.")

doc.add_paragraph()
h("3. Backend Work (Light)", 15, (20, 20, 20))
bullet("Removed the college-email restriction so users can register and sign in with any personal email.")
bullet("Full REST API already in place: authentication, skills, matches, exchanges, messaging, ratings, and leaderboard.")
bullet("Seeded demo data (50 skills, 14 users) for instant testing; all 14 end-to-end tests passing.")

doc.add_paragraph()
h("4. Result", 15, (20, 20, 20))
para(
    "The app now delivers a distinctive, premium UI with a clean black aesthetic, smooth motion, "
    "and a glass navigation dock, while the backend continues to reliably power login, matching, "
    "exchanges, chat, and ratings."
)

out = r"C:\Users\Admin\skillswap-mvp\SkillSwap_Progress_Report.docx"
doc.save(out)
print("Saved:", out)